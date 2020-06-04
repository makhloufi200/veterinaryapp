from django.shortcuts import render
from .models import Animals
from django.shortcuts import render, redirect, get_object_or_404
from django.http import Http404, HttpResponse
from .forms import AnimalsForm
import csv
import docx
# Create your views here.

def export_animal_csv(request):
	objs = Animals.objects.all()
	#response['Content-Disposition'] = 'attachement; filename="animal.csv"'
	#writer = csv.writer(response)
	#writer.writerow(['name'])
	animals = Animals.objects.all().values_list('name')
	doc = docx.Document()
	p = doc.add_paragraph('List of Animals')
		
	
	#doc.save('multipleParagraphs.docx')
	row_count = len(animals)
	table = doc.add_table(rows=row_count, cols=1)
	
	for i in range(row_count):
		row = table.rows[i]
		
		#for animal in animals:
		row.cells[0].text = animals[i]
	doc.save('multipleParagraphs.docx')	
	return render(request, 'animal/list_animals.html', {'animals': objs})
	
def list_animal(request):
	objs = Animals.objects.all()
	
	return render(request, 'animal/list_animals.html', {'animals': objs})
	
def datail_animal(request, animal_id):
    animal = get_object_or_404(Animals, pk=animal_id)
    return render(request, 'animal/animal.html', {'animal': animal})	
	
def create_animal(request):
    if not request.user.is_authenticated:
        raise Http404

    form = AnimalsForm(request.POST or None)
    if form.is_valid():
        instance = form.save(commit=False)
        instance.save()
        return redirect('list_animal')

    return render(request, 'animal/animal_form.html', {'form': form})
	
def update_animal(request, animal_id):
	if not request.user.is_authenticated:
		raise Http404
	instance = get_object_or_404(Animals, pk=animal_id)

	form = AnimalsForm(request.POST or None, instance=instance)
	if form.is_valid():
		instance = form.save(commit=False)
		instance.save()
		return redirect('list_animal')

	return render(request, 'animal/animal_form.html', {"form": form,"instance": instance})


def animal_delete(request, animal_id):
    instance = get_object_or_404(Animals, pk=animal_id)
    instance.delete()

    return redirect("list_animal")													